import argparse
import sys
import os
import threading
import subprocess

from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QPushButton, QLabel, QLineEdit, QVBoxLayout, \
    QHBoxLayout, QComboBox

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document


def create_label_entry(label_text, entry_widget):  # 创建一个水平布局，使得标签和输入部件水平
    layout = QHBoxLayout()
    label = QLabel(label_text)
    layout.addWidget(label)
    layout.addWidget(entry_widget)
    return layout


def search_file(path, keyword):  # 用来暴力搜索文件
    match_files = []

    for root, dirstory, files in os.walk(path):
        for file_search in files:
            if keyword in file_search:
                match_files.append(os.path.join(root, file_search))

    return match_files


class Window(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("迅捷搜索")
        self.setGeometry(100, 100, 450, 150)

        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        self.EntryInputKeyword = QLineEdit(self)
        self.EntryInputPath = QLineEdit(self)
        self.EntryInputPath.setText(os.getcwd())  # 默认路径为当前工作目录

        layout.addLayout(create_label_entry("关键字", self.EntryInputKeyword))
        layout.addLayout(create_label_entry("路径", self.EntryInputPath))

        self.select = QComboBox()
        self.select.addItems(["PDF", "Word", "MD"])
        layout.addLayout(create_label_entry("导出格式", self.select))

        self.buttonExport = QPushButton("导出", self)
        self.buttonQuit = QPushButton("退出", self)

        # 连接功能
        self.buttonExport.clicked.connect(self.export_the_file)
        self.buttonQuit.clicked.connect(self.close)

        button_layout = QHBoxLayout()
        button_layout.addStretch()
        button_layout.addWidget(self.buttonExport)
        button_layout.addWidget(self.buttonQuit)

        layout.addLayout(button_layout)

        self.LabelOutputinstruct = QLabel("该路径下的文件为：")
        self.LabelOutputmessage = QLabel()
        layout.addWidget(self.LabelOutputinstruct)
        layout.addWidget(self.LabelOutputmessage)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        self.EntryInputKeyword.textChanged.connect(self.display_the_file)

    def display_the_file(self):  # 用来显示文件的
        self.LabelOutputinstruct.setText("正在搜索文件，请稍候...")
        # 使用线程执行搜索过程
        search_thread = threading.Thread(target=self.updateUI_after_display)
        search_thread.start()

    def updateUI_after_display(self):
        self.LabelOutputinstruct.setText("该路径下的文件为：")
        match_files = search_file(self.EntryInputPath.text(), self.EntryInputKeyword.text())

        if match_files:
            self.LabelOutputmessage.setText("\n".join(match_files))
        else:
            self.LabelOutputmessage.setText("未找到匹配文件")

    def export_the_file(self):  # 用来导出文件
        self.LabelOutputinstruct.setText("正在导出文件，请稍候...")
        # 使用线程执行搜索过程
        search_thread = threading.Thread(target=self.updateUI_after_export)
        search_thread.start()

    def updateUI_after_export(self):
        flag = self.select.currentText()
        match_files = search_file(self.EntryInputPath.text(), self.EntryInputKeyword.text())

        if flag == 'PDF':
            c = canvas.Canvas("output.pdf", pagesize=letter)
            y = 700
            for file_search in match_files:
                c.drawString(100, y, file_search)
                y -= 20
            c.save()
            subprocess.Popen(["start", "output.pdf"], shell=True)  # 打开导出的pdf文件
        elif flag == 'Word':
            doc = Document()
            for file_search in match_files:
                doc.add_paragraph(file_search)
            doc.save("output.docx")
            subprocess.Popen(["start", "output.docx"], shell=True)  # 打开导出的word文件
        elif flag == 'MD':
            with open("output.md", "w") as file:
                for file_search in match_files:
                    file.write("- " + file_search + "\n")
            subprocess.Popen(["start", "output.md"], shell=True)  # 打开导出的md文件
        else:
            print("导出文件错误")

        self.LabelOutputinstruct.setText("导出文件成功")


def main():
    parser = argparse.ArgumentParser(description="迅捷搜索程序")
    parser.add_argument("--search", type=str, help="按关键字搜索文件")
    parser.add_argument("--export", type=str, choices=["word", "pdf", "md"], help="将搜索结果导出为指定格式")
    parser.add_argument("--path", type=str, help="选择路径")
    args = parser.parse_args()

    if args.path:
        os.chdir(args.path)

    if args.search:
        # 如果提供了 --search 参数，则执行搜索操作
        search_keyword = args.search
        search_path = args.path if args.path else os.getcwd()
        match_files = search_file(search_path, search_keyword)
        if match_files:
            print("匹配的文件：")
            for file in match_files:
                print(file)
        else:
            print("未找到匹配文件")

    if args.export:
        # 如果提供了 --export 参数，则执行导出操作
        export_format = args.export
        search_keyword = args.search
        search_path = args.path if args.path else os.getcwd()
        match_files = search_file(search_path, search_keyword)

        if export_format == 'pdf':
            c = canvas.Canvas("output.pdf", pagesize=letter)
            y = 700
            for file_search in match_files:
                c.drawString(100, y, file_search)
                y -= 20
            c.save()
            subprocess.Popen(["start", "output.pdf"], shell=True)  # 打开导出的pdf文件
        elif export_format == 'word':
            doc = Document()
            for file_search in match_files:
                doc.add_paragraph(file_search)
            doc.save("output.docx")
            subprocess.Popen(["start", "output.docx"], shell=True)  # 打开导出的word文件
        elif export_format == 'md':
            with open("output.md", "w") as file:
                for file_search in match_files:
                    file.write("- " + file_search + "\n")
            subprocess.Popen(["start", "output.md"], shell=True)  # 打开导出的md文件
        else:
            print("导出文件错误")

    else:
        print("nothing")


if __name__ == "__main__":
    if len(sys.argv) == 1:
        # 如果没有命令行参数，则以PyCharm运行方式运行
        app = QApplication(sys.argv)
        window = Window()
        window.show()
        sys.exit(app.exec_())
    else:
        # 如果存在命令行参数，则以命令行方式运行
        main()


